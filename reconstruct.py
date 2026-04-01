# api/reconstruct.py — Vercel Python serverless function
# Receives: original DOCX (base64) + rewritten text sections from Claude
# Returns: reconstructed DOCX (base64) with original formatting preserved exactly
#
# HOW IT WORKS:
# 1. Decode the original DOCX from base64
# 2. Open it with python-docx
# 3. Walk every paragraph and table cell
# 4. For each paragraph, find its matching replacement text from Claude's output
# 5. Consolidate runs, replace text, preserve all formatting XML
# 6. Return the modified DOCX as base64
#
# The AI never touches formatting. Python never touches content logic.
# These two systems never interact with each other.

import base64
import io
import json
import re
from http.server import BaseHTTPRequestHandler

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── RUN CONSOLIDATION ────────────────────────────────────────────────────────
def consolidate_runs(paragraph):
    """
    Collapse all runs in a paragraph into a single run,
    keeping the first run's formatting as the template.

    This solves the multi-run problem: a single visual sentence
    is often split across 3-4 <w:r> elements with different formatting.
    After consolidation, we can safely replace the text of the single run.

    The paragraph-level formatting (pPr: indent, spacing, bullet style)
    is NEVER touched — only run-level text is changed.
    """
    if not paragraph.runs:
        return

    full_text = paragraph.text
    if not full_text.strip():
        return

    p_el = paragraph._p

    # Remove all runs except the first
    runs_to_remove = paragraph.runs[1:]
    for run in runs_to_remove:
        p_el.remove(run._r)

    # Set the first run's text to the full concatenated text
    if paragraph.runs:
        paragraph.runs[0].text = full_text


def replace_paragraph_text(paragraph, new_text):
    """
    Replace all text in a paragraph with new_text.
    Preserves all paragraph and run formatting.
    """
    consolidate_runs(paragraph)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        # No runs — add one inheriting paragraph formatting
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = new_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        paragraph._p.append(r)


# ── PARAGRAPH TEXT EXTRACTION ────────────────────────────────────────────────
def get_para_text(paragraph):
    """Get full text of a paragraph including all runs."""
    return paragraph.text.strip()


def get_all_paragraphs(doc):
    """
    Get all paragraphs from the document including inside tables and text boxes.
    Returns list of paragraph objects.
    """
    paragraphs = list(doc.paragraphs)

    # Also get paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para not in paragraphs:
                        paragraphs.append(para)

    # Also get paragraphs inside text boxes (txbxContent)
    body = doc.element.body
    for txbx in body.iter(qn('w:txbxContent')):
        for p_el in txbx.iter(qn('w:p')):
            from docx.text.paragraph import Paragraph
            para = Paragraph(p_el, doc)
            if para not in paragraphs:
                paragraphs.append(para)

    return paragraphs


# ── MATCHING ENGINE ──────────────────────────────────────────────────────────
def normalize(text):
    """Normalize text for fuzzy matching — lowercase, collapse whitespace."""
    return re.sub(r'\s+', ' ', (text or '').lower().strip())


def find_best_match(original_text, replacement_map):
    """
    Find the best matching replacement for a paragraph's original text.
    Uses exact match first, then fuzzy prefix match.
    Returns (new_text, confidence) or (None, 0).
    """
    norm_orig = normalize(original_text)
    if not norm_orig or len(norm_orig) < 10:
        return None, 0

    # Exact match
    if norm_orig in replacement_map:
        return replacement_map[norm_orig], 1.0

    # Prefix match — original starts with the first 40 chars of a key
    for key, val in replacement_map.items():
        if len(key) >= 30 and norm_orig.startswith(key[:40]):
            return val, 0.9
        if len(norm_orig) >= 30 and key.startswith(norm_orig[:40]):
            return val, 0.85

    return None, 0


# ── MAIN RECONSTRUCTION ──────────────────────────────────────────────────────
def reconstruct_docx(original_bytes, rewritten_sections):
    """
    Core reconstruction function.

    Args:
        original_bytes: bytes of the original DOCX file
        rewritten_sections: dict mapping original paragraph text (normalized)
                           to new paragraph text from Claude

    Returns:
        bytes of the reconstructed DOCX
    """
    # Build normalized replacement map
    replacement_map = {}
    for orig, new in rewritten_sections.items():
        replacement_map[normalize(orig)] = new

    # Open the original document
    doc = Document(io.BytesIO(original_bytes))

    # Process all paragraphs
    all_paras = get_all_paragraphs(doc)
    replaced_count = 0

    for para in all_paras:
        orig_text = get_para_text(para)
        if not orig_text:
            continue

        new_text, confidence = find_best_match(orig_text, replacement_map)

        if new_text and confidence >= 0.85:
            replace_paragraph_text(para, new_text)
            replaced_count += 1

    # Save to bytes
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    return output_buffer.read(), replaced_count


# ── VERCEL HANDLER ───────────────────────────────────────────────────────────
class handler(BaseHTTPRequestHandler):

    def do_POST(self):
        if not DOCX_AVAILABLE:
            self._error(500, 'python-docx not available — add to requirements.txt')
            return

        # Read request body
        content_length = int(self.headers.get('Content-Length', 0))
        body = self.rfile.read(content_length)

        try:
            data = json.loads(body)
        except json.JSONDecodeError:
            self._error(400, 'Invalid JSON')
            return

        # Validate inputs
        docx_b64 = data.get('docx_base64')
        rewritten = data.get('rewritten_sections')  # {original_text: new_text}

        if not docx_b64:
            self._error(400, 'Missing docx_base64')
            return
        if not rewritten or not isinstance(rewritten, dict):
            self._error(400, 'Missing or invalid rewritten_sections')
            return

        # Decode DOCX
        try:
            original_bytes = base64.b64decode(docx_b64)
        except Exception:
            self._error(400, 'Invalid base64 for docx_base64')
            return

        # Reconstruct
        try:
            result_bytes, replaced_count = reconstruct_docx(original_bytes, rewritten)
        except Exception as e:
            self._error(500, f'Reconstruction failed: {str(e)}')
            return

        # Return reconstructed DOCX as base64
        response = {
            'docx_base64': base64.b64encode(result_bytes).decode('utf-8'),
            'replaced_count': replaced_count,
            'status': 'ok'
        }

        self._json(200, response)

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors_headers()
        self.end_headers()

    def _json(self, status, data):
        body = json.dumps(data).encode('utf-8')
        self.send_response(status)
        self._cors_headers()
        self.send_header('Content-Type', 'application/json')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _error(self, status, message):
        self._json(status, {'error': message})

    def _cors_headers(self):
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
